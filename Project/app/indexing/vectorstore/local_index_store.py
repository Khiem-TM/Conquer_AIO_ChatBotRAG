from __future__ import annotations

import hashlib
import json
import re
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.indexing.config import settings
from app.indexing.schemas import IndexOperationResult, IndexStatus
from app.shared.utils import get_logger

logger = get_logger(__name__)


@dataclass(slots=True)
class SourceDocument:
    source_id: str
    source_name: str
    file_path: Path
    relative_path: str
    updated_at_ns: int
    file_size: int
    text: str
    source_hash: str


class LocalIndexStore:
    """Local store helper for reading files, chunking, and JSON snapshot export.

    JSON is kept only as a lightweight local export/debug artifact. The primary
    local source of truth should be SQLite via ``SQLiteIndexStore``.
    """

    def __init__(self) -> None:
        self._index_data: dict[str, Any] | None = None
        self._project_dir = Path(__file__).resolve().parents[3]

    def load_index_data(self) -> dict[str, Any] | None:
        if self._index_data is not None:
            return self._index_data
        storage_path = self.get_storage_path()
        if not storage_path.exists():
            return None
        self._index_data = json.loads(storage_path.read_text(encoding='utf-8'))
        return self._index_data

    def write_index_data(self, index_data: dict[str, Any]) -> None:
        storage_path = self.get_storage_path()
        storage_path.parent.mkdir(parents=True, exist_ok=True)
        storage_path.write_text(json.dumps(index_data, ensure_ascii=False, indent=2), encoding='utf-8')
        self._index_data = index_data

    def scan_source_files(self) -> list[Path]:
        data_input_dir = self.get_data_input_dir()
        if not data_input_dir.exists():
            logger.warning('data_input directory does not exist: %s', data_input_dir)
            return []

        supported_extensions = {'.md', '.txt', '.docx', '.pdf'}
        return [
            file_path
            for file_path in sorted(data_input_dir.rglob('*'))
            if file_path.is_file()
            and file_path.suffix.lower() in supported_extensions
            and not file_path.name.startswith('~$')
            and not file_path.name.startswith('.')
        ]

    def collect_documents(self, source_files: list[Path]) -> list[SourceDocument]:
        documents: list[SourceDocument] = []
        for file_path in source_files:
            text = self.read_source_text(file_path)
            if not text:
                continue
            relative_path = file_path.relative_to(self.get_data_input_dir()).as_posix()
            documents.append(
                SourceDocument(
                    source_id=self.build_source_id(file_path),
                    source_name=file_path.name,
                    file_path=file_path,
                    relative_path=relative_path,
                    updated_at_ns=file_path.stat().st_mtime_ns,
                    file_size=file_path.stat().st_size,
                    text=text,
                    source_hash=self.compute_hash(text),
                )
            )
        return documents

    def prepare_chunk_records(self, documents: list[SourceDocument]) -> tuple[list[dict[str, Any]], list[str]]:
        raw_chunks: list[dict[str, Any]] = []
        texts: list[str] = []

        for document in documents:
            chunks = self.split_text(document.text)
            for index, chunk_text in enumerate(chunks, start=1):
                cleaned_text = chunk_text.strip()
                if not cleaned_text:
                    continue
                chunk_hash = self.compute_hash(f'{document.source_hash}:{cleaned_text}')
                metadata = {
                    'source_name': document.source_name,
                    'relative_path': document.relative_path,
                    'file_path': document.relative_path,
                    'source_hash': document.source_hash,
                    'chunk_hash': chunk_hash,
                    'chunk_index': index,
                    'char_count': len(cleaned_text),
                }
                raw_chunks.append(
                    {
                        'source_id': document.source_id,
                        'source_name': document.source_name,
                        'chunk_id': f'{document.source_id}_chunk_{index:03d}',
                        'chunk_index': index,
                        'text': cleaned_text,
                        'chunk_hash': chunk_hash,
                        'metadata': metadata,
                    }
                )
                texts.append(cleaned_text)

        return raw_chunks, texts

    def build_sources_payload(self, documents: list[SourceDocument]) -> dict[str, dict[str, Any]]:
        payload: dict[str, dict[str, Any]] = {}
        for document in documents:
            payload[document.source_id] = {
                'source_name': document.source_name,
                'file_path': document.relative_path,
                'relative_path': document.relative_path,
                'updated_at_ns': document.updated_at_ns,
                'file_size': document.file_size,
                'source_hash': document.source_hash,
                'document_char_count': len(document.text),
                'chunk_count': 0,
            }
        return payload

    def build_status_response(self, index_data: dict[str, Any] | None) -> IndexStatus:
        if not index_data:
            return IndexStatus(embedding_model=settings.embedding_model)

        return IndexStatus(
            ready=bool(index_data.get('chunks')),
            embedding_backend=index_data.get('embedding_backend', 'pending'),
            embedding_model=index_data.get('embedding_model', settings.embedding_model),
            total_sources=len(index_data.get('sources', {})),
            total_chunks=len(index_data.get('chunks', [])),
            built_at=index_data.get('built_at'),
        )

    def build_operation_response(
        self,
        index_data: dict[str, Any],
        message: str,
        latency_ms: int,
        updated_sources: list[str],
        deleted_sources: list[str],
    ) -> IndexOperationResult:
        status_response = self.build_status_response(index_data)
        return IndexOperationResult(
            ready=status_response.ready,
            embedding_backend=status_response.embedding_backend,
            embedding_model=status_response.embedding_model,
            total_sources=status_response.total_sources,
            total_chunks=status_response.total_chunks,
            built_at=status_response.built_at,
            message=message,
            latency_ms=latency_ms,
            updated_sources=updated_sources,
            deleted_sources=deleted_sources,
        )

    def empty_index_data(self) -> dict[str, Any]:
        return {
            'schema_version': settings.index_schema_version,
            'chunker_version': settings.index_chunker_version,
            'embedding_backend': 'pending',
            'embedding_model': settings.embedding_model,
            'built_at': None,
            'sources': {},
            'chunks': [],
        }

    def build_source_id(self, file_path: Path) -> str:
        relative_path = file_path.relative_to(self.get_data_input_dir()).as_posix()
        return relative_path.replace('/', '__')

    def read_source_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == '.docx':
            return self._read_docx_text(file_path)
        if suffix == '.pdf':
            return self._read_pdf_text(file_path)
        try:
            return file_path.read_text(encoding='utf-8').strip()
        except UnicodeDecodeError:
            return file_path.read_text(encoding='utf-8', errors='ignore').strip()

    def _read_docx_text(self, file_path: Path) -> str:
        try:
            from docx import Document  # type: ignore[import]
            from docx.table import Table  # type: ignore[import]
            from docx.text.paragraph import Paragraph  # type: ignore[import]
        except ModuleNotFoundError:
            logger.warning('python-docx not installed, falling back to raw XML parse for %s', file_path)
            return self._read_docx_text_fallback(file_path)

        try:
            doc = Document(str(file_path))
        except Exception as exc:
            logger.warning('Failed to open DOCX %s: %s', file_path, exc)
            return ''

        blocks: list[str] = []

        def _iter_block_items(parent):  # type: ignore[no-untyped-def]
            """Yield paragraphs and tables in document order."""
            from docx.oxml.ns import qn  # type: ignore[import]
            from docx.oxml import OxmlElement  # type: ignore[import]
            for child in parent.element.body:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'p':
                    yield Paragraph(child, parent)
                elif tag == 'tbl':
                    yield Table(child, parent)

        try:
            for item in _iter_block_items(doc):
                if isinstance(item, Paragraph):
                    text = item.text.strip()
                    if text:
                        blocks.append(text)
                elif isinstance(item, Table):
                    for row in item.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_cells:
                            blocks.append(' | '.join(row_cells))
        except Exception as exc:
            logger.warning('Error iterating DOCX body %s: %s', file_path, exc)
            # Best-effort: fall back to paragraphs only
            blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        return '\n'.join(blocks).strip()

    def _read_docx_text_fallback(self, file_path: Path) -> str:
        """Legacy raw-XML fallback when python-docx is unavailable."""
        try:
            with zipfile.ZipFile(file_path) as archive:
                xml_bytes = archive.read('word/document.xml')
        except Exception as exc:
            logger.warning('Failed to read DOCX source %s: %s', file_path, exc)
            return ''
        try:
            root = ET.fromstring(xml_bytes)
        except ET.ParseError as exc:
            logger.warning('Failed to parse DOCX XML %s: %s', file_path, exc)
            return ''
        text_nodes: list[str] = []
        for node in root.iter():
            if node.tag.endswith('}t') and node.text:
                text_nodes.append(node.text)
        return '\n'.join(text_nodes).strip()

    def _read_pdf_text(self, file_path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ModuleNotFoundError:
            logger.warning('pypdf is not installed; skip PDF source %s', file_path)
            return ''

        try:
            reader = PdfReader(str(file_path))
            page_texts = [(page.extract_text() or '').strip() for page in reader.pages]
            return '\n\n'.join(text for text in page_texts if text).strip()
        except Exception as exc:
            logger.warning('Failed to read PDF source %s: %s', file_path, exc)
            return ''

    def split_text(self, text: str) -> list[str]:
        blocks = [block.strip() for block in re.split(r'\n\s*\n', text) if block.strip()]
        chunks: list[str] = []
        current_chunk = ''

        for block in blocks:
            next_chunk = f'{current_chunk}\n\n{block}'.strip() if current_chunk else block
            if len(next_chunk) <= settings.index_chunk_size:
                current_chunk = next_chunk
                continue

            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ''

            if len(block) <= settings.index_chunk_size:
                current_chunk = block
                continue

            chunks.extend(self.split_long_block(block))

        if current_chunk:
            chunks.append(current_chunk)
        return chunks

    def split_long_block(self, text: str) -> list[str]:
        chunks: list[str] = []
        start = 0
        chunk_size = max(300, settings.index_chunk_size)
        overlap = max(0, min(settings.index_chunk_overlap, chunk_size // 2))

        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(text):
                break
            start = max(start + 1, end - overlap)
        return chunks

    def compute_hash(self, text: str) -> str:
        return hashlib.sha256(text.encode('utf-8')).hexdigest()

    def get_data_input_dir(self) -> Path:
        return self.resolve_path(settings.index_data_input_dir)

    def get_storage_path(self) -> Path:
        return self.resolve_path(settings.index_storage_path)

    def utc_now(self) -> str:
        return datetime.now(timezone.utc).isoformat()

    def resolve_path(self, value: str) -> Path:
        path = Path(value)
        if path.is_absolute():
            return path
        return (self._project_dir / path).resolve()
